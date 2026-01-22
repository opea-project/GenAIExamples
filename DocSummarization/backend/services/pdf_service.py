"""
Document Extraction Service
Handles PDF and DOCX text extraction with OCR support for image-based PDFs
"""

from pypdf import PdfReader
from docx import Document
from pdf2image import convert_from_path
import pytesseract
from typing import Optional
import logging
import os
import config

logger = logging.getLogger(__name__)


class PDFService:
    """Service for extracting text from PDF and DOCX files"""

    def __init__(self):
        """Initialize document extraction service"""
        logger.info("Document Extraction Service initialized")

    def extract_text(self, file_path: str) -> str:
        """
        Extract text from PDF or DOCX file

        Args:
            file_path: Path to document file

        Returns:
            Extracted text content

        Raises:
            Exception: If document extraction fails
        """
        try:
            filename_lower = file_path.lower()

            if filename_lower.endswith('.pdf'):
                return self._extract_from_pdf(file_path)
            elif filename_lower.endswith(('.docx', '.doc')):
                return self._extract_from_docx(file_path)
            else:
                raise Exception(f"Unsupported file type: {file_path}")

        except Exception as e:
            logger.error(f"Document extraction error: {str(e)}")
            raise Exception(f"Failed to extract text from document: {str(e)}")

    def _extract_from_pdf(self, pdf_path: str) -> str:
        """
        Extract text from PDF file with automatic OCR fallback for image-based PDFs
        Implements page limit to prevent processing extremely large PDFs

        Args:
            pdf_path: Path to PDF file

        Returns:
            Extracted text content
        """
        logger.info(f"Extracting text from PDF: {pdf_path}")

        text_content = ""

        # First, try standard text extraction
        with open(pdf_path, "rb") as file:
            pdf_reader = PdfReader(file)
            num_pages = len(pdf_reader.pages)

            logger.info(f"PDF has {num_pages} pages")

            # Apply page limit
            max_pages = config.MAX_PDF_PAGES
            pages_to_process = min(num_pages, max_pages)

            if num_pages > max_pages:
                logger.warning(f"PDF has {num_pages} pages. Processing only first {max_pages} pages to prevent timeout.")
                text_content += f"[Note: This PDF has {num_pages} pages. Processing first {max_pages} pages only.]\n\n"

            for page_num, page in enumerate(pdf_reader.pages[:pages_to_process], 1):
                page_text = page.extract_text()
                text_content += page_text + "\n"
                logger.debug(f"Extracted {len(page_text)} characters from page {page_num}")

        extracted_length = len(text_content.strip())
        logger.info(f"Extracted {extracted_length} characters from PDF")

        # If no text was extracted, the PDF is likely image-based - use OCR
        if extracted_length < 50:  # Threshold for considering PDF as image-based
            logger.info("PDF appears to be image-based or has minimal text. Using OCR...")
            text_content = self._extract_with_ocr(pdf_path, max_pages=pages_to_process)
            extracted_length = len(text_content.strip())
            logger.info(f"OCR extracted {extracted_length} characters from PDF")

        return text_content.strip()

    def _extract_with_ocr(self, pdf_path: str, max_pages: Optional[int] = None) -> str:
        """
        Extract text from PDF using OCR (for image-based PDFs)

        Args:
            pdf_path: Path to PDF file
            max_pages: Maximum number of pages to process (None = all pages)

        Returns:
            Extracted text using OCR
        """
        try:
            logger.info(f"Starting OCR extraction for: {pdf_path}")

            # Convert PDF pages to images
            images = convert_from_path(pdf_path, dpi=300)
            total_pages = len(images)
            logger.info(f"Converted PDF to {total_pages} images")

            # Apply page limit if specified
            if max_pages and total_pages > max_pages:
                logger.warning(f"OCR: Processing only first {max_pages} of {total_pages} pages")
                images = images[:max_pages]

            text_content = ""
            pages_to_process = len(images)

            # Perform OCR on each page
            for page_num, image in enumerate(images, 1):
                logger.info(f"Running OCR on page {page_num}/{pages_to_process}")
                page_text = pytesseract.image_to_string(image)
                text_content += page_text + "\n"
                logger.debug(f"OCR extracted {len(page_text)} characters from page {page_num}")

            extracted_length = len(text_content.strip())
            logger.info(f"OCR successfully extracted {extracted_length} characters")

            return text_content.strip()

        except Exception as e:
            logger.error(f"OCR extraction failed: {str(e)}")
            raise Exception(f"Failed to extract text using OCR: {str(e)}")

    def _extract_from_docx(self, docx_path: str) -> str:
        """
        Extract text from DOCX file

        Args:
            docx_path: Path to DOCX file

        Returns:
            Extracted text content
        """
        logger.info(f"Extracting text from DOCX: {docx_path}")

        text_content = ""

        doc = Document(docx_path)

        # Extract text from all paragraphs
        for paragraph in doc.paragraphs:
            text_content += paragraph.text + "\n"

        # Extract text from tables if present
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text_content += cell.text + " "
                text_content += "\n"

        extracted_length = len(text_content.strip())
        logger.info(f"Successfully extracted {extracted_length} characters from DOCX")

        return text_content.strip()

    def get_pdf_metadata(self, pdf_path: str) -> dict:
        """
        Get PDF metadata (title, author, pages, etc.)

        Args:
            pdf_path: Path to PDF file

        Returns:
            Dictionary with PDF metadata
        """
        try:
            with open(pdf_path, "rb") as file:
                pdf_reader = PdfReader(file)

                metadata = {
                    "num_pages": len(pdf_reader.pages),
                    "title": pdf_reader.metadata.get("/Title", "Unknown") if pdf_reader.metadata else "Unknown",
                    "author": pdf_reader.metadata.get("/Author", "Unknown") if pdf_reader.metadata else "Unknown",
                }

                return metadata

        except Exception as e:
            logger.error(f"Error getting PDF metadata: {str(e)}")
            return {}


# Global PDF service instance
pdf_service = PDFService()
