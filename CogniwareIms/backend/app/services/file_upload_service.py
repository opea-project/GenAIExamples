# Copyright (C) 2024 Intel Corporation
# SPDX-License-Identifier: Apache-2.0
"""File Upload Service Handles xlsx, csv, pdf, docx file uploads and processing for knowledge base Optimized for Intel
Xeon processors."""

import hashlib
import logging
import os
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional

import pandas as pd

# File processing libraries
try:
    from openpyxl import load_workbook
except ImportError:
    load_workbook = None

try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

try:
    from docx import Document
except ImportError:
    Document = None

from .embedding_service import embedding_service
from .knowledge_manager import knowledge_manager
from .retrieval_service import retrieval_service

logger = logging.getLogger(__name__)


class FileUploadService:
    """
    Handles file uploads and processing for knowledge base
    Supports: CSV, XLSX, PDF, DOCX
    """

    SUPPORTED_EXTENSIONS = {".csv", ".xlsx", ".pdf", ".docx"}
    MAX_FILE_SIZE = 100 * 1024 * 1024  # 100MB

    def __init__(self, upload_dir: str = "/app/uploads"):
        self.upload_dir = Path(upload_dir)
        self.upload_dir.mkdir(exist_ok=True, parents=True)

        # Create subdirectories by type
        for ext in self.SUPPORTED_EXTENSIONS:
            (self.upload_dir / ext[1:]).mkdir(exist_ok=True, parents=True)

    def validate_file(self, filename: str, file_size: int) -> tuple[bool, str]:
        """Validate uploaded file.

        Returns:
            (is_valid, error_message)
        """
        # Check size
        if file_size > self.MAX_FILE_SIZE:
            return (
                False,
                f"File too large. Maximum size: {self.MAX_FILE_SIZE // (1024*1024)}MB",
            )

        # Check extension
        ext = Path(filename).suffix.lower()
        if ext not in self.SUPPORTED_EXTENSIONS:
            return (
                False,
                f"Unsupported file type. Supported: {', '.join(self.SUPPORTED_EXTENSIONS)}",
            )

        return True, ""

    async def save_file(self, filename: str, content: bytes) -> Path:
        """Save uploaded file to disk.

        Returns:
            Path to saved file
        """
        # Generate unique filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        file_hash = hashlib.sha256(content).hexdigest()[:8]
        ext = Path(filename).suffix.lower()
        base_name = Path(filename).stem

        unique_filename = f"{base_name}_{timestamp}_{file_hash}{ext}"

        # Save to appropriate subdirectory
        file_path = self.upload_dir / ext[1:] / unique_filename

        with open(file_path, "wb") as f:
            f.write(content)

        logger.info(f"Saved file: {file_path}")
        return file_path

    async def process_csv(self, file_path: Path) -> Dict[str, Any]:
        """Process CSV file and add to knowledge base."""
        try:
            df = pd.read_csv(file_path)

            documents = []
            for idx, row in df.iterrows():
                # Create text representation
                text_parts = [f"{col}: {row[col]}" for col in df.columns if pd.notna(row[col])]
                text = " | ".join(text_parts)

                documents.append(
                    {
                        "text": text,
                        "metadata": {
                            "source": "csv_upload",
                            "filename": file_path.name,
                            "row_index": idx,
                            "uploaded_at": datetime.now().isoformat(),
                            "file_type": "csv",
                        },
                    }
                )

            # Add to knowledge base
            result = await knowledge_manager.add_knowledge_batch(
                documents=documents, source=f"csv_upload_{file_path.stem}"
            )

            return {
                "success": True,
                "file_type": "csv",
                "filename": file_path.name,
                "rows_processed": len(documents),
                "documents_added": result.get("added", 0),
                "total_documents": result.get("total_documents", 0),
            }

        except Exception as e:
            logger.error(f"CSV processing error: {e}")
            return {"success": False, "error": str(e), "file_type": "csv"}

    async def process_xlsx(self, file_path: Path) -> Dict[str, Any]:
        """Process XLSX file and add to knowledge base."""
        try:
            if load_workbook is None:
                return {
                    "success": False,
                    "error": "openpyxl not installed. Run: pip install openpyxl",
                }

            # Read all sheets
            excel_file = pd.ExcelFile(file_path)
            all_documents = []

            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)

                for idx, row in df.iterrows():
                    text_parts = [f"{col}: {row[col]}" for col in df.columns if pd.notna(row[col])]
                    text = " | ".join(text_parts)

                    all_documents.append(
                        {
                            "text": text,
                            "metadata": {
                                "source": "xlsx_upload",
                                "filename": file_path.name,
                                "sheet": sheet_name,
                                "row_index": idx,
                                "uploaded_at": datetime.now().isoformat(),
                                "file_type": "xlsx",
                            },
                        }
                    )

            # Add to knowledge base
            result = await knowledge_manager.add_knowledge_batch(
                documents=all_documents, source=f"xlsx_upload_{file_path.stem}"
            )

            return {
                "success": True,
                "file_type": "xlsx",
                "filename": file_path.name,
                "sheets_processed": len(excel_file.sheet_names),
                "rows_processed": len(all_documents),
                "documents_added": result.get("added", 0),
                "total_documents": result.get("total_documents", 0),
            }

        except Exception as e:
            logger.error(f"XLSX processing error: {e}")
            return {"success": False, "error": str(e), "file_type": "xlsx"}

    async def process_pdf(self, file_path: Path) -> Dict[str, Any]:
        """Process PDF file and add to knowledge base."""
        try:
            if PdfReader is None:
                return {
                    "success": False,
                    "error": "pypdf not installed. Run: pip install pypdf>=4.0.0",
                }

            documents = []

            with open(file_path, "rb") as file:
                pdf_reader = PdfReader(file)
                total_pages = len(pdf_reader.pages)

                for page_num in range(total_pages):
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()

                    if text.strip():  # Only add non-empty pages
                        documents.append(
                            {
                                "text": text,
                                "metadata": {
                                    "source": "pdf_upload",
                                    "filename": file_path.name,
                                    "page": page_num + 1,
                                    "total_pages": total_pages,
                                    "uploaded_at": datetime.now().isoformat(),
                                    "file_type": "pdf",
                                },
                            }
                        )

            # Add to knowledge base
            result = await knowledge_manager.add_knowledge_batch(
                documents=documents, source=f"pdf_upload_{file_path.stem}"
            )

            return {
                "success": True,
                "file_type": "pdf",
                "filename": file_path.name,
                "pages_processed": len(documents),
                "documents_added": result.get("added", 0),
                "total_documents": result.get("total_documents", 0),
            }

        except Exception as e:
            logger.error(f"PDF processing error: {e}")
            return {"success": False, "error": str(e), "file_type": "pdf"}

    async def process_docx(self, file_path: Path) -> Dict[str, Any]:
        """Process DOCX file and add to knowledge base."""
        try:
            if Document is None:
                return {
                    "success": False,
                    "error": "python-docx not installed. Run: pip install python-docx",
                }

            doc = Document(file_path)
            documents = []

            # Process paragraphs
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)

            # Split into chunks (every 5 paragraphs or ~500 words)
            chunk_size = 5
            for i in range(0, len(full_text), chunk_size):
                chunk = " ".join(full_text[i : i + chunk_size])

                if chunk.strip():
                    documents.append(
                        {
                            "text": chunk,
                            "metadata": {
                                "source": "docx_upload",
                                "filename": file_path.name,
                                "chunk": i // chunk_size + 1,
                                "uploaded_at": datetime.now().isoformat(),
                                "file_type": "docx",
                            },
                        }
                    )

            # Process tables
            for table_idx, table in enumerate(doc.tables):
                table_text = []
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells])
                    table_text.append(row_text)

                if table_text:
                    documents.append(
                        {
                            "text": "\n".join(table_text),
                            "metadata": {
                                "source": "docx_upload",
                                "filename": file_path.name,
                                "table": table_idx + 1,
                                "uploaded_at": datetime.now().isoformat(),
                                "file_type": "docx",
                            },
                        }
                    )

            # Add to knowledge base
            result = await knowledge_manager.add_knowledge_batch(
                documents=documents, source=f"docx_upload_{file_path.stem}"
            )

            return {
                "success": True,
                "file_type": "docx",
                "filename": file_path.name,
                "chunks_processed": len(documents),
                "documents_added": result.get("added", 0),
                "total_documents": result.get("total_documents", 0),
            }

        except Exception as e:
            logger.error(f"DOCX processing error: {e}")
            return {"success": False, "error": str(e), "file_type": "docx"}

    async def process_file(self, file_path: Path) -> Dict[str, Any]:
        """Route file to appropriate processor based on extension."""
        ext = file_path.suffix.lower()

        processors = {
            ".csv": self.process_csv,
            ".xlsx": self.process_xlsx,
            ".pdf": self.process_pdf,
            ".docx": self.process_docx,
        }

        processor = processors.get(ext)
        if processor is None:
            return {"success": False, "error": f"No processor for file type: {ext}"}

        return await processor(file_path)

    async def upload_and_process(self, filename: str, content: bytes) -> Dict[str, Any]:
        """Complete upload and processing workflow.

        Args:
            filename: Original filename
            content: File bytes

        Returns:
            Processing result with statistics
        """
        # Validate
        is_valid, error_msg = self.validate_file(filename, len(content))
        if not is_valid:
            return {"success": False, "error": error_msg}

        # Save file
        file_path = await self.save_file(filename, content)

        # Process file
        result = await self.process_file(file_path)

        # Add file path to result
        result["file_path"] = str(file_path)

        return result

    def list_uploaded_files(self, file_type: Optional[str] = None) -> List[Dict[str, Any]]:
        """List all uploaded files."""
        files = []

        search_dirs = []
        if file_type:
            search_dirs.append(self.upload_dir / file_type)
        else:
            for ext in self.SUPPORTED_EXTENSIONS:
                search_dirs.append(self.upload_dir / ext[1:])

        for search_dir in search_dirs:
            if search_dir.exists():
                for file_path in search_dir.iterdir():
                    if file_path.is_file():
                        stat = file_path.stat()
                        files.append(
                            {
                                "filename": file_path.name,
                                "type": file_path.suffix[1:],
                                "size": stat.st_size,
                                "uploaded_at": datetime.fromtimestamp(stat.st_ctime).isoformat(),
                                "path": str(file_path),
                            }
                        )

        return sorted(files, key=lambda x: x["uploaded_at"], reverse=True)


# Global instance
file_upload_service = FileUploadService(upload_dir=os.getenv("UPLOAD_DIR", "/app/uploads"))
